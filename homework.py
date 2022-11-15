#安装包 pip install python-docx 
#国内镜像站： -i https://pypi.tuna.tsinghua.edu.cn/simple
import turtle,time #绘制单段数码管
from PIL import Image

from docx.document import Document
from docx.text.paragraph import Paragraph #段落处理
from docx.shared import RGBColor  #设置颜色
from docx import Document
from docx.shared import Pt  #设置字号，间距

def drawLine(draw):
    turtle.pendown()if draw else turtle.penup()
    turtle.fd(40)
    turtle.right(90)
def drawDigit(digit):#根据数字绘制七段数码管
    drawLine(True) if digit in [2,3,4,5,6,8,9] else drawLine(False)
    drawLine(True)if digit in [0,1,3,4,5,6,7,8,9] else drawLine(False)
    drawLine( True)if digit in [0,2,3,5,6,8,9] else drawLine(False)
    drawLine(True) if digit in [0,2,6,8] else drawLine(False)
    turtle.left(90)
    drawLine(True) if digit in [0,4,5,6,8,9] else drawLine(False)
    drawLine(True) if digit in [0,2,3,5,6,7,8,9] else drawLine(False)
    drawLine(True) if digit in [0,1,2,3,4,7,8,9] else drawLine(False)
    turtle.left(180)
    turtle.penup()#为绘制后续数字确定位置
    turtle.fd(20) # 为绘制后续数字确定位置
def drawDate(date):
    turtle.pencolor("red")
    for i in date:
         if i == '-':
            turtle.write('年', font=("Arial", 18,"normal"))
            turtle.pencolor("green")
            turtle.fd(40)
         elif i == '=':
            turtle.write('月', font=("Arial", 18,"normal"))
            turtle.pencolor("blue")
            turtle.fd(40)
         elif i == '+':
            turtle.write('日', font=("Arial", 18,"normal"))
         else:
            drawDigit(eval(i))

document = Document()#新建文档
number = document.add_paragraph().add_run('xxx') #添加段落及段内文字
number.font.size= Pt(26) #设置文字字号
number.font.color.rgb=RGBColor(128, 0, 128) #设置文字颜色，紫色rgb
#下同
college = document.add_paragraph().add_run('xxx')
college.font.size = Pt(26)
college.font.color.rgb=RGBColor(128, 0, 128)

paragraph_3= document.add_paragraph()
name=paragraph_3.add_run('xxx')
name.font.size = Pt(26)
name.font.color.rgb=RGBColor(128, 0, 128)

paragraph_format = paragraph_3.paragraph_format #仅能处理段落，故paragraph_3不直接add，先设置个段落出来
paragraph_format.space_after = Pt(45) #60像素以windows默认dpi：96转换为45pt，详见：https://www.cnblogs.com/zhenzhong/p/3348567.html

turtle.setup(800,350,200,200)
turtle.penup()
turtle.fd(-300)
turtle.pensize(5)
drawDate(time.strftime('%Y-%m=%d+',time.gmtime()))
turtle.hideturtle()
# 保存图片
ts = turtle.getscreen()
ts.getcanvas().postscript(file="date.eps")

date=document.add_picture('date.jpeg')

document.save('homework.docx') #文档保存